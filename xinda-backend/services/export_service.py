from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
from PIL import Image

def format_shanghai_time(dt):
    if dt is None:
        return "未知"
    if dt.tzinfo is None:
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    return dt.strftime("%Y-%m-%d %H:%M:%S (Asia/Shanghai)")

class ExportService:
    def export_to_word(self, record):
        doc = Document()
        
        doc.add_heading(f'文档处理结果: {record.original_filename}', 0)
        
        doc.add_paragraph(f'处理时间: {format_shanghai_time(record.upload_time)}')
        doc.add_paragraph(f'文件类型: {record.file_type}')
        
        doc.add_heading('提取的日文文本', level=1)
        if record.ocr_text:
            doc.add_paragraph(record.ocr_text)
        else:
            doc.add_paragraph('(无OCR文本)')
        
        doc.add_heading('中文翻译', level=1)
        if record.translated_text:
            doc.add_paragraph(record.translated_text)
        else:
            doc.add_paragraph('(无翻译文本)')
        
        if record.file_type == 'jpg':
            doc.add_heading('原始图像', level=1)
            try:
                doc.add_picture(record.file_path, width=Inches(6))
            except Exception:
                doc.add_paragraph('(图像加载失败)')
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    
    def export_ocr_only(self, record):
        doc = Document()
        doc.add_heading(f'日文识别稿: {record.original_filename}', 0)
        doc.add_paragraph(f'处理时间: {format_shanghai_time(record.upload_time)}')
        doc.add_paragraph(f'文件类型: {record.file_type}')
        
        doc.add_heading('提取的日文文本', level=1)
        if record.ocr_text:
            doc.add_paragraph(record.ocr_text)
        else:
            doc.add_paragraph('(无OCR文本)')
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    
    def export_translate_only(self, record):
        doc = Document()
        doc.add_heading(f'中文翻译稿: {record.original_filename}', 0)
        doc.add_paragraph(f'处理时间: {format_shanghai_time(record.upload_time)}')
        doc.add_paragraph(f'文件类型: {record.file_type}')
        
        doc.add_heading('中文翻译', level=1)
        if record.translated_text:
            doc.add_paragraph(record.translated_text)
        else:
            doc.add_paragraph('(无翻译文本)')
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
